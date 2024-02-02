import re
import time
from datetime import date
from io import BytesIO
import json

from django import forms
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import Document
from docx.shared import Pt

from django.contrib.auth.decorators import login_required, user_passes_test, permission_required
from django.utils.decorators import method_decorator
from django.contrib.auth.models import Group, User
from django.shortcuts import render
from django.db import transaction
from django.db.models import Q, Sum
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from dateutil.relativedelta import relativedelta

from app_worktime.forms import WorkTimeForm, WorkTime_rowForm, WorkTimeRepForm
from app_worktime.models import *
from app_worktime.dovidnyk import Dovidnyk

import matplotlib
import matplotlib.pyplot as plt
import io
import urllib
import base64
import numpy as np

matplotlib.use('Agg')  # Вибрати не-GUI фоновий режим


# -----  Головна сторінка сайту  ---
def main_page_view(request):
    template_name = 'main_page.html'
    data = {"Title": "Головна сторінка", "permit": request.user.get_all_permissions()}
    return render(request, template_name, context=data)


# -----  Сканер QR  ---
def scanner_qr(request):
    template_name = 'scanner.html'
    data = {"Title": "Сканер QR-коду", "permit": request.user.get_all_permissions()}
    return render(request, template_name, context=data)


@login_required
def login_main_page_view(request):
    template_name = 'main_page.html'
    data = {"Title": "Головна сторінка", "permit": request.user.get_all_permissions()}
    return render(request, template_name, context=data)


# =============== Довідники ===================
# --- Підрозділи
departments = Dovidnyk(dov_model=Department, dov_prefix="dep")

# --- Співробітники
employees = Dovidnyk(dov_model=Employee, dov_prefix="empl")


# ======== >>>

class QRcode:
    """ ---------------  Обєкт для роботи з QR-кодом -------------------
    """
    def __init__(self):
        pass

    def make_qrcode(self, params):
        # ********************* формування QR-code
        import qrcode

        # URL-рядок з параметрами
        url_with_params = params
        # QR-код
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(url_with_params)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        # Збережіть або відправте картинку до клієнта
        # img.save("static/qrcode.png")  # Збереження QR-коду в файл

        return img

    def read_qrcode(self):
        pass

class Paginator:
    """ ---------------  Обєкт для керування Пагінатором Веб-сторінок -------------------
        параметри в рядку:
           options = {"page": 3, "limit": 12, "period": true, "work_date": "2023-09-01", "tab_nom": "1523"}
    """

    def __init__(self, options=''):
        self.page_limit = 8
        self.element_remove = ['total_pages', 'total_row', 'pages']
        if options:
            self.params = json.loads(options)
            self.params['page'] = int(self.params.get('page', 1))
            self.params['limit'] = int(self.params.get('limit', self.page_limit))
        else:
            self.params = {'page': 1, 'limit': self.page_limit, 'total_pages': 1}

    def set_total_page(self, total_page=0, total_row=0):
        self.params['total_pages'] = total_page
        self.params['total_row'] = total_row  # всього рядків
        # формуємо список сторінок
        if self.params['total_pages']:
            self.params['pages'] = [i for i in range(1, self.params['total_pages'] + 1)]
        else:
            self.params['pages'] = []
        # якщо номер поточної сторінки більше total_page, то поточна сторінка = total_page
        if self.params['page'] > self.params['total_pages']:  self.params['page'] = self.params['total_pages']

    def get_options_list(self, dop_dict: dict):
        # - формуємо словник з Пагинатора та додаткового словника
        if self.params:
            res_dict = self.params.copy()
            # додаткові параметри, якщо є
            if dop_dict:
                res_dict.update(dop_dict)
            # видаляємо елементи, які не потрібні
            for el in self.element_remove:
                res_dict.pop(el, None)
            options_list = json.dumps(res_dict, separators=(',', ':'))
        else:
            options_list = ''

        return options_list


class WorkTime:
    """ ---------------  Облік робочого часу НТЦ -------------------
    Права доступа
    --------------
    Користувачі Джанго(User) мають логін = Таб.номер
    Співробітники (Employee) мають поле Таб.номер
    Після входу в систему для юзера, який має доступ до  застосунку app_worktime
    шукаємо відповідий запис в Employee.
    Далі, юзер буде мати доступ тільки до тих записів таблиці AccountWorkTime, для яких Employee=Цей юзер
      !! Для юзера:  tab_nom=Таб.номер, deps=None
      Юзер входить до Групи 'Облік робочого часу'

    ------------
    Начальник бюро, начальник відділу крім цього матиме доступ до перегляду даних по своїм співробітникам.
    Керівництво НТЦ матиме доступ до перегляду всіх даних

      !! Для керівництва:   tab_nom="", deps=[перелік бюро та нових підрозділів]
      Керівник входить до Групи 'Облік робочого часу(керівники)'

    Якщо користувач входить до групи  'Облік робочого часу(дирекція)'
                можна дивитися Звіти в цілому по НТЦ

            from django.contrib.auth.models import User
            usr = User.objects.get(username='your_username')
    """

    def __init__(self, direction="1", work_template="work_time.html", report_template="work_time_report.html"):
        self.direction = direction  # ('1', 'НТЦ'), ('2', 'Технологи')
        # параметри для різних дирекцій
        title_dirs = {'1': 'НТЦ', '2': 'служб Головного інженера'}
        prefix_dirs = {'1': 'NTC', '2': 'TECH'}
        self.title_dir = title_dirs[self.direction]
        self.prefix_dir = prefix_dirs[self.direction]
        self.rep_code = f"0{self.direction}"
        self.num_days_after_month = 4

        self.work_template: str = work_template  # шаблон сторінки розподілу робочого часу НТЦ
        self.report_template: str = report_template  # шаблон сторінки розподілу робочого часу НТЦ
        self.img_pie_chart = ''  # кругова діаграма в кінці Звітів
        self.total_table_html = []

        self.docx_templ_1 = f"work_time_{self.prefix_dir}_dot.docx"
        self.docx_templ_23 = f"work_time_svod_{self.prefix_dir}_dot.docx"

        self.work_filters_template = "worktime_filters.html"

    def has_user_perm(self, request):
        user = request.user
        dep_permit = []
        tab_nom = ''
        # Перевіряємо доступ поточного користувача до app_worktime  (група Облік робочого часу НТЦ)
        # can_change = 1 if 'app_worktime.view_accountworktime' in user.get_all_permissions() else 0
        can_change = 1 if user.has_perm(f"app_worktime.{self.prefix_dir}_worktime") else 0
        # <--- якщо доступ є
        if can_change == 1:
            #  user = admin
            if user.is_superuser:
                dep_permit = None
                tab_nom = "all"
            else:
                #  доступ до конкретного користувача
                # якщо ім'я - це таб.номер
                if user.username.strip().isdigit():
                    tabn = user.username
                else:
                    # шукаємо таб.номер в полі Прізвище  003 - Михайлик
                    tabn = user.last_name.strip().split('-')[0]
                # знаходимо співробітника по таб.номеру з User
                empl = Employee.objects.filter(tab_nom=tabn.strip()).first()
                if empl: tab_nom = empl.tab_nom
            # ----->
        else:
            # << --- Перевіряємо доступ поточного користувача до app_worktime (група Облік робочого часу(керівники)
            # if 'app_worktime.change_accountworktime' in user.get_all_permissions():
            if user.has_perm(f"app_worktime.{self.prefix_dir}_worktime_headman"):
                can_change = 1
                #  доступ до підрозділів НТЦ( )
                for group in Group.objects.filter(user=user):
                    if group.name.startswith('dir-') or group.name.startswith('Облік'):
                        continue
                    else:
                        dep_permit.append(group.name.split()[0])
                # --- Перевіряємо доступ поточного користувача до app_worktime (група Облік робочого часу(дирекція)
                #     можна дивитися Звіти в цілому по НТЦ
                # if 'app_worktime.view_reportworktimetempl' in user.get_all_permissions():
                if user.has_perm(f"app_worktime.{self.prefix_dir}_worktime_dir"):
                    can_change = 10
            # ---->>

        return can_change, tab_nom, dep_permit

    # @method_decorator(user_passes_test(lambda user: user.has_perm('app_worktime.change_accountworktime')))
    # @method_decorator(permission_required('app_worktime.change_accountworktime', raise_exception=True), name='dispatch')
    @method_decorator(login_required)
    def list_worktime_filters(self, request, options=''):
        """ ---- Виводимо список робіт по заданим фільтрам ----
            worktime_filters/{"period":true,"work_date":"2023-10-02","tab_nom":"1277","filters":{"zakaz":"ВЗ0001"}}
            options = {"period":true,"work_date":"2023-10-02","tab_nom":"1277", "main_field": "zakaz", "filters": {"zakaz":"ВЗ0001"} }
            filters = {"zakaz": "ВЗ0001", "employee": 10}
        """
        #  Параметри
        if not options: return HttpResponse("Не задані критерії відбору інформації !")
        #  розбираємо параметри
        params = json.loads(options) if options else {}
        if params:
            filters = params.get('filters')
        else:
            filters = {}
        if not filters: return HttpResponse("Не задані параметри відбору інформації !")

        # ------ отримуємо список робіт по фільтрам
        queryset = AccountWorkTime.objects.filter(**filters).order_by("employee", "work_date")
        #  Підготовка даних для таблиці
        dict_zak = get_list_zak()  # назви Заказів
        fields_name = ['work_date', 'zakaz', 'employee', 'work_content', 'work_detail', 'doc_kind', 'doc_designation',
                       'work_time', 'work_percent', 'id']
        #  заголовки стовпців таблиці
        title_fields = get_list_field_title(AccountWorkTime, fields_name)
        title_fields[0] = "Дата\nроботи"
        title_fields[1] = "Виробниче\nЗамовлення"
        title_fields[8] = "% виконання\nдокумента"
        #  Якщо є головне поле в фільтрі, по якому формується звіт (наприклад - Заказ)
        main_field = params.get('main_field')
        if main_field:
            #  для виводу в заголовок
            first_row = queryset.first()
            if first_row:
                name_main_field = getattr(first_row, main_field)
                if main_field == 'zakaz':
                    name_main_field = f"{name_main_field}   {dict_zak.get(name_main_field, '???')}"  # назва заказу
                    name_main_field = f' по проекту "{name_main_field}" '
            else:
                name_main_field = '?????'
            # видаляємо це поле з таблиці
            if main_field in fields_name:
                index = fields_name.index(main_field)
                del fields_name[index]
                del title_fields[index]
        else:
            name_main_field = ''

        #  --- список рядків для видачі в таблицю
        total_time = 0
        rows_list_value = []
        for row in queryset:
            values = {}
            for f in fields_name:
                values[f] = getattr(row, f)
                if isinstance(values[f], date):  values[f] = values[f].strftime('%d.%m.%Y')
                if f == 'employee':
                    values['employee'] = row.employee.full_name
                elif f == 'zakaz':
                    values['nam_zak'] = dict_zak.get(values['zakaz'], '???')
            rows_list_value.append(values)
            total_time += getattr(row, 'work_time', 0)  # сума витраченого часу по списку робіт
        #  якщо в таблиці є заказ - додаємо стовпчик з назвою заказу
        if 'zakaz' in fields_name:
            title_fields.insert(2, "Назва Замовлення")
            fields_name.insert(2, 'nam_zak')
        # формування рядка таблиці для виводу суми витраченого часу
        total_time_row = {f: '' for f in fields_name}
        total_time_row['work_time'] = '***' + str(total_time)
        rows_list_value.insert(0, total_time_row)

        # --- Пагінатор зроблено на клієнті, задаємо початкові параметри
        params = {'limit': 15, 'total_row': len(rows_list_value), 'page': 1}

        title = f'Розподіл робочого часу співробітників {self.title_dir} {name_main_field}'
        options_list = ''

        data = {"Title": title, "permit": request.user.get_all_permissions(),
                "title_fields": title_fields, "rows_list_value": rows_list_value,
                "params": params, "options_list": options_list,
                }
        return render(request, self.work_filters_template, context=data)

    # -------------------------------------
    @method_decorator(login_required)
    def list_worktime(self, request, period=True, work_date=None, tab_nom=None, edit_mode=0, id_row=0, options=''):
        """ ===============================================================================
            ****      Сеанс перегляду та коригування розподілу робочого часу НТЦ    *******
            ===============================================================================
            Параметри запуску:
            period - період:
                якщо +, то work_date - перший день місяця та відбір даних по місяцю
                якщо - , то  work_date = робочий день в періоді та відбір по конкретній даті
            work_date - дата початку місяця, або дата робочого дня
            tab_nom - визначаеться з групи юзера

            Тільки режим перегляду :
               вибір періода, Дати та Співробітника,тформування таблиці розподілу робочого часу на основі
               збережених даних з моделі AccountWorkTime,
               працює форма  userform = WorkTimeForm
           Для коригування рядка викликаємо метод edit_worktime - нова сторінка

            edit_mode==0 - Режим перегляду :
                вибір періода, Дати та Співробітника,тформування таблиці розподілу робочого часу на основі
                збережених даних з моделі AccountWorkTime,
                працює форма  userform = WorkTimeForm
            edit_mode==1 - Режим коригування :
                таблиця розподілу робочого часуь сформована по Даті та Співробітнику, зміна цих параметрів заблокована
                обираємо рядок для коригування, він передається в форму WorkNime_rowForm
                дані записуються в модель AccountWorkTime
            edit_mode==3 - Копіювання рядка в межах поточного місяця
            edit_mode==4 - Копіювання рядка на наступний місяць
            -------- передача параметрів в рядку
            options = {"page": 3, "limit": 12, "period": true, "work_date": "2023-09-01", "tab_nom": "1523"}
        """
        #  <---- перевірка дозволів
        can_change, tabnom_permit, dep_permit = self.has_user_perm(request)
        # доступу немає
        if not (can_change and (tabnom_permit or dep_permit)):
            return HttpResponse("Ви не маєте доступу до інформації !")

        if tabnom_permit != "all" and tab_nom and not dep_permit and tab_nom != tabnom_permit:
            return HttpResponse(f"Ви не маєте доступу до табельного номера {tab_nom} !")
        if not tab_nom: tab_nom = tabnom_permit
        # перевірка можливості коригування по періоду
        can_edit = True
        can_copy_next_month = True
        # ---->

        # ********************* формування QR-code
        # if options:
        #     import qrcode
        #
        #     # URL-рядок з параметрами
        #     url_with_params = request.path
        #     # QR-код
        #     qr = qrcode.QRCode(
        #         version=1,
        #         error_correction=qrcode.constants.ERROR_CORRECT_L,
        #         box_size=10,
        #         border=4,
        #     )
        #     qr.add_data(url_with_params)
        #     qr.make(fit=True)
        #     img = qr.make_image(fill_color="black", back_color="white")
        #     # Збережіть або відправте картинку до клієнта
        #     img.save("static/qrcode.png")  # Збереження QR-коду в файл

        # --- Пагінатор та параметри
        paginator = Paginator(options)
        if options:
            period = paginator.params.get('period', True)
            work_date = paginator.params.get('work_date')
            tab_nom = paginator.params.get('tab_nom', '')

        if isinstance(period, str): period = True if period == 'True' else False
        if isinstance(work_date, str): work_date = datetime.strptime(work_date, "%Y-%m-%d")
        if work_date and period: work_date = work_date.replace(day=1)
        # якщо це перенос на наступний місяць - збільшуємо дату на місяць
        if edit_mode == 4:   work_date = work_date + relativedelta(months=1)

        #  розрахункові змінні
        # total_pages = 0
        row_form = title_fields = rows_list_value = total_time_row = None
        list_doc_kind = doc_kind_for_work_content = form_value = None

        # --- Відбір даних по заданих параметрах
        if request.method == "POST":
            userform = WorkTimeForm(request.POST, tab_nom=tab_nom, dep_permit=dep_permit)
            # <<----- перевірки даних форми параметрів відбору
            if not userform.is_valid():
                return HttpResponse("Не пройшла перевірка даних, відправлених у Формі !!!")
            form_value = userform.cleaned_data  # отримуємо очищені дані
            # - повнота даних
            if not (form_value['par_work_date'] and form_value['employee']):
                return HttpResponse("Не повна інформація у відправленій Формі !!!")

            #  ---- параметри відбору
            period = form_value['period']
            work_date = form_value['par_work_date'].replace(day=1) if period else form_value['par_work_date']
            tab_nom = form_value['employee'].tab_nom

        #  ------------ ініціація форми параметрів---------
        userform = WorkTimeForm(tab_nom=tab_nom, dep_permit=dep_permit)
        if tab_nom and tab_nom != "all":
            userform.initial['employee'] = Employee.objects.get(tab_nom=tab_nom)

        if request.method == "POST" or (work_date and tab_nom):
            userform.initial['par_work_date'] = work_date.strftime('%Y-%m-%d')

            if request.method == "POST":
                userform.initial['period'] = form_value['period']
                userform.initial['employee'] = form_value['employee']
            #  --- повернення після редагування або видалення рядка
            elif work_date and tab_nom:
                userform.initial['period'] = period
                userform.initial['employee'] = Employee.objects.get(tab_nom=tab_nom)

            # --- якщо таблицю нарахувань сформовано - форма параметрів в режимі перегляду
            userform.fields['par_work_date'].widget.attrs.update({'readonly': 'readonly'})
            userform.fields['period'].widget.attrs.update({'disabled': True})

            # --- перевірка можливості коригування по періоду:
            # не дозволяємо зміни, якщо поточна дата більша за num_days_after_month після початку наступного місяця,
            # або якщо місяць поточної дати більше місяця work_date на 2+
            cur_date = datetime.now().date()
            if cur_date.month - work_date.month > 1:
                can_edit = can_copy_next_month = False
            elif cur_date.month - work_date.month == 1:
                can_edit = cur_date.day < self.num_days_after_month

        #  --- запуск пустої форми
        else:
            pass

        selected_par = work_date and tab_nom
        #  ---------- ТАБЛИЦЯ РОЗПОДІЛУ РОБОЧОГО ЧАСУ ПО СПІВРОБІТНИКУ --------------
        if edit_mode == 0:
            # <<--------- Підготовка даних для виводу в таблицю:  форма, заголовки, поля таблиці  -----
            #  рядки
            fields_name = ['work_date', 'zakaz', 'work_content', 'work_detail', 'doc_kind', 'doc_designation',
                           'work_time', 'work_percent', 'id']
            title_fields = get_list_field_title(AccountWorkTime, fields_name)
            title_fields[0] = "Дата\nроботи"
            title_fields[1] = "Виробниче\nЗамовлення"
            title_fields[7] = "% виконання\nдокумента"

            dict_zak = get_list_zak()

            rows_list_value = []
            total_time = 0
            # --- якщо вибрані параметри ---
            if selected_par:
                # --- рядки таблиці
                if period:
                    work_date_end = work_date + relativedelta(months=1) - relativedelta(days=1)
                    queryset = AccountWorkTime.objects.filter(
                        Q(employee__tab_nom=tab_nom) & Q(work_date__gte=work_date) & Q(work_date__lte=work_date_end))
                else:
                    queryset = AccountWorkTime.objects.filter(Q(employee__tab_nom=tab_nom) & Q(work_date=work_date))

                # ----- Paginator ---
                # кількість елементів в queryset
                count = queryset.count()
                # кількість сторінок
                paginator.set_total_page(count // paginator.params['limit'] + (1 if count % paginator.params['limit'] else 0), count)
                # останній порядковий номер номер запису сторінки <page>
                last_row = paginator.params['page'] * paginator.params['limit']
                # перший порядковий номер номер запису сторінки <page>
                first_row = last_row - paginator.params['limit'] + 1

                #  --- список рядків для видачі в таблицю
                current_row = 1
                for row in queryset:
                    #  якщо це рядки вибраної сторінки
                    if not paginator.params['page'] or (first_row <= current_row <= last_row):
                        values = {}
                        for f in fields_name:
                            values[f] = getattr(row, f)
                            if isinstance(values[f], date):  values[f] = values[f].strftime('%d.%m.%Y')
                            if f == 'zakaz':  values['nam_zak'] = dict_zak.get(values['zakaz'], '???')
                        rows_list_value.append(values)
                        total_time += getattr(row, 'work_time', 0)  # сума витраченого часу
                    current_row += 1

            title_fields.insert(2, "Назва Замовлення")
            fields_name.insert(2, 'nam_zak')
            # формування рядка таблиці для виводу суми витраченого часу
            if total_time: total_time_row = [total_time if f == 'work_time' else '' for f in fields_name if f != 'id']

        elif edit_mode:
            #  ---------- БЛАНК КОРИГУВАННЯ РОЗПОДІЛУ РОБОЧОГО ЧАСУ ПО СПІВРОБІТНИКУ --------------
            row_acc = AccountWorkTime.objects.filter(pk=id_row).first()
            row_form = WorkTime_rowForm(instance=row_acc)
            row_form.fields['zakaz'].choices = get_list_zak(1)  # для отримання актуального списку замовлень в формі !!!
            row_form.fields['work_content'].queryset = WorkContent.objects.filter(direction=self.direction)
            if self.direction == '1':
                row_form.fields['kol_normoper'].widget = forms.HiddenInput()

            #  якщо новий запис - співробітника беремо з форми переметрів
            if not row_acc:
                row_form.initial['employee'] = userform.initial['employee']

            # якщо вибрана конкретна дата, то фіксуємо її
            if not period:
                if not row_acc or edit_mode == 3: row_form.initial['work_date'] = work_date.strftime('%Y-%m-%d')
                row_form.fields['work_date'].widget.attrs.update({'readonly': 'readonly'})

            # якщо вибраний період - фіксуємо межі місяця
            else:
                date_start = work_date.strftime('%Y-%m-%d')
                date_end = work_date.date() + relativedelta(months=1) - relativedelta(days=1)
                if date_end > date.today(): date_end = date.today()
                date_end = date_end.strftime('%Y-%m-%d')
                row_form.fields['work_date'].widget.attrs.update({'min': date_start, 'max': date_end})

            # якщо це копіювання рядка
            if edit_mode in [3, 4]:
                #  - обнуляємо id
                id_row = 0
                # якщо це перенос на наступний місяць - коригуємо дату операції відповідно до періоду
                if edit_mode == 4 and date_start:
                    row_form.initial['work_date'] = date_start
                edit_mode = 1

            # ---- Підготовка списків Видів документів для кожного Змісту роботи
            #       далі буде перемикати Javascript на сторінці
            # всі види документів
            list_doc_kind = DocKind.objects.values_list("id", "npp", "name")
            # для кожного Змісту роботи формуємо перелік доступних видів документів - в рядок через кому
            doc_kind_for_work_content = ['']
            # for wk_cont in WorkContent.objects.filter(direction=self.direction):
            for wk_cont in WorkContent.objects.all():
                arr_str = ''
                for arr in wk_cont.doc_kinds.values_list("id", flat=True):
                    arr_str += f"{arr}" if arr_str == '' else f",{arr}"
                doc_kind_for_work_content.insert(wk_cont.id, arr_str)

            # -------------------------------->>>

        title = f"Розподіл робочого часу співробітників {self.title_dir} по проектах за місяць"
        if work_date:
            title = title.replace("місяць", month_ua[work_date.month])
        #  --- Пагінатор для виводу списку робіт
        if selected_par:
            options_list = paginator.get_options_list({'period': period, 'work_date': work_date.strftime("%Y-%m-%d"), 'tab_nom': tab_nom})
        else:
            options_list = ''

        data = {"Title": title, "permit": request.user.get_all_permissions(), "edit_mode": edit_mode,
                "can_edit": can_edit, "can_copy_next_month": can_copy_next_month,
                "selected_par": selected_par, "period": period,
                "work_date": work_date.strftime('%Y-%m-%d') if work_date else '',
                # "work_date_next": (work_date + relativedelta(months=1)).strftime('%Y-%m-%d') if work_date else '',
                "tab_nom": tab_nom, "id_row": id_row,
                "form": userform, "row_form": row_form,
                "title_fields": title_fields, "rows_list_value": rows_list_value, "total_time_row": total_time_row,
                "list_doc_kind": list_doc_kind, "doc_kind_for_work_content": doc_kind_for_work_content,
                "params": paginator.params, "options_list": options_list,
                }
        return render(request, self.work_template, context=data)

    @method_decorator(login_required)
    # @permission_required('app_worktime.TECH_worktime', raise_exception=True)
    def get_report(self, request, kind_rep=1, mode='', options=''):
        """ ===========================================================================
        ------ Підсумкові звіти по виконаним роботам ------
        mode=='print' -  вигрузка у Word
        mode=='chart' -  виведення діаграми
        ==============================================================================="""
        #  <---- перевірка дозволів
        can_change, tabnom_permit, dep_permit = self.has_user_perm(request)
        # доступу немає
        if not (can_change and (tabnom_permit or dep_permit)):
            return HttpResponse("Ви не маєте доступу до інформації !")
        tab_nom = tabnom_permit
        # ---->

        # --- параметри
        if isinstance(kind_rep, str): kind_rep = int(kind_rep)
        work_date = work_date_end = None
        title_fields = report = []
        worker = self.title_dir
        title = ''
        period = False
        self.img_pie_chart = ''
        self.total_table_html = []
        execution_time_str = ''
        start_time = 0
        params = {}
        filters = {}
        options_list = ''

        # --- Формування Звіту по заданих параметрах
        if request.method == "POST" or options:
            #  --- час формування звіту----
            start_time = time.time()
            # -- запуск форми
            userform = WorkTimeRepForm(request.POST, tab_nom=tab_nom, dep_permit=dep_permit)
            # <<----- перевірки даних форми параметрів відбору
            if not userform.is_valid():
                return HttpResponse("Не пройшла перевірка даних, відправлених у Формі !!!")
            form_value = userform.cleaned_data  # отримуємо очищені дані
            # - повнота даних
            if not (form_value['par_work_date'] and (form_value['par_work_date_end'] or not form_value['par_work_date'])
                    and
                    (form_value['employee'] or form_value['dep'] or form_value['ntc_total'] or form_value['section'])):
                return HttpResponse("Не повна інформація у відправленій Формі !!!")

            if tabnom_permit != "all" and not dep_permit and form_value['employee'].tab_nom != tabnom_permit:
                return HttpResponse(f"Ви не маєте доступу до табельного номера {form_value['employee'].tab_nom} !")

            #  ---- параметри відбору
            period = form_value['period']
            #  дати початку періоду та  кінця періоду
            if kind_rep in (3, 4):
                work_date = form_value['par_work_date'].replace(day=1)
                work_date_end = form_value['par_work_date_end'].replace(day=1)
                work_date_end = work_date_end + relativedelta(months=1) - relativedelta(days=1)
            else:
                if not period:
                    work_date = form_value['par_work_date'].replace(day=1)
                    work_date_end = work_date + relativedelta(months=1) - relativedelta(days=1)
                else:
                    work_date = form_value['par_work_date']
                    work_date_end = form_value['par_work_date_end']

            employee = form_value['employee']
            section = form_value['section']
            dep = form_value['dep']
            ntc_total = form_value['ntc_total']

            # --- ініціація форми ---
            userform = WorkTimeRepForm(tab_nom=tab_nom, dep_permit=dep_permit)
            userform.initial['par_work_date'] = work_date.strftime('%Y-%m-%d')
            if kind_rep == 3:
                userform.initial['period'] = True
            else:
                userform.initial['period'] = period
            userform.initial['par_work_date_end'] = work_date_end.strftime('%Y-%m-%d')
            #  --- відбір по людина/бюро/відділ/НТЦ
            if employee:
                userform.initial['employee'] = employee
                worker = f"{employee} "
            elif section:
                userform.initial['section'] = section
                worker = f" '{section.dep_name}' " if 'бюро' in section.dep_name.lower() else f"бюро '{section.dep_name}' "
            elif dep:
                userform.initial['dep'] = dep
                worker = f"підрозділу '{dep.dep_name}' "
            elif ntc_total:
                userform.initial['ntc_total'] = ntc_total
                worker = f"по {self.title_dir} "

            # <<--------- Підготовка даних для виводу в таблицю:  форма, заголовки, поля таблиці  -----
            rep_code = self.rep_code
            #    -- шаблон звіту
            templ_rep = ReportWorkTimeTempl.objects.filter(rep_code=rep_code)
            #    ------ Розподіл робочого часу по заказах - фільтр по періоду
            filters.update({'work_date__gte': work_date.isoformat(), 'work_date__lte': work_date_end.isoformat()})
            queryset = AccountWorkTime.objects.select_related('doc_kind__doctype').filter(
                Q(work_date__gte=work_date) & Q(work_date__lte=work_date_end))
            #    -- фільтрація в такому порядку:   таб.ном / бюро/ підрозділ / НТЦ
            if employee:
                # фільтр по співробітнику
                filters.update({"employee": employee.id})
                queryset = queryset.filter(employee=employee)
            elif section:
                #  фільтр по бюро
                filters.update({"employee__section": section.id})
                queryset = queryset.filter(employee__section=section)
            elif dep:
                #  фільтр по підрозділу
                filters.update({"employee__section__dep_code__startswith": dep.dep_code[:2]})
                queryset = queryset.filter(employee__section__dep_code__startswith=dep.dep_code[:2])
            elif ntc_total and can_change == 10:
                #  фільтр по НТЦ
                filters.update({"employee__section__dep_direction": '1'})
                queryset = queryset.filter(employee__section__dep_direction='1')
            else:
                #  помилка - жодного параметру відбору не передано
                queryset = None

            #  --- час формування звіту ----
            end_time = time.time()
            execution_time_str += f"Підготовка даних - {(end_time - start_time):.2f} сек\n"
            start_time = end_time

            # ---------------  формування даних звіту на базі шаблону
            #  ==============  Підсумковий звіт по виконаним роботам за місяць ============
            if kind_rep == 1:
                # === Звіт --
                title_fields, report = self.get_report_01(templ_rep, queryset)
                #         --- час формування звіту ----
                end_time = time.time()
                execution_time_str += f"Формування таблиць - {(end_time - start_time):.2f} сек\n"
                start_time = end_time
                # === Діаграма --
                self.get_chart_12(queryset)
                #         --- час формування звіту ----
                end_time = time.time()
                execution_time_str += f"Формування діаграми - {(end_time - start_time):.2f} сек\n"
                start_time = end_time
                # ---
            elif kind_rep == 2:
                # якщо формуємо звіт за місяць - додаємо Фактично відпрацьовано по табелю
                if period:
                    fact_time_list = {}
                else:
                    fact_time_list = {row.tab_nom: row.fact_time for row in WorkTimeFact.objects.filter(work_date=work_date)}
                # === Звіт --
                title_fields, report, params, time_rep_detail = self.get_report_02(templ_rep, queryset, fact_time_list)
                #         --- час формування звіту ----
                execution_time_str += time_rep_detail
                end_time = time.time()
                execution_time_str += f"Формування таблиць - {(end_time - start_time):.2f} сек\n"
                start_time = end_time
                # === Діаграма --
                self.get_chart_12(queryset)
                #         --- час формування звіту ----
                end_time = time.time()
                execution_time_str += f"Формування діаграми - {(end_time - start_time):.2f} сек\n"
                start_time = end_time
                # ---
            elif kind_rep == 3:
                title_fields, report = self.get_report_03(templ_rep, queryset, work_date, work_date_end)
                #         --- час формування звіту ----
                end_time = time.time()
                execution_time_str += f"Формування таблиць - {(end_time - start_time):.2f} сек\n"
                start_time = end_time
                # ---
            elif kind_rep == 4:
                title_fields, report, params = self.get_report_04(templ_rep, queryset, options)
                #  ---- додавання параметрів фільтрації для деталізації по замовленням
                params_details: dict = params.copy()
                for row in report:
                    filters['zakaz'] = re.match(r'^\S+', row[1]).group(0)
                    params_details['filters'] = filters
                    options_details = json.dumps(params_details, separators=(',', ':'))
                    row.append(options_details)
                # ----------------
                #         --- час формування звіту ----
                end_time = time.time()
                execution_time_str += f"Формування таблиць - {(end_time - start_time):.2f} сек\n"
                start_time = end_time
                # ---

        # ------ запуск пустої форми
        else:
            userform = WorkTimeRepForm(tab_nom=tab_nom, dep_permit=dep_permit)
            userform.initial['par_work_date_end'] = date.today().strftime('%Y-%m-%d')

        # ----  Ініціалізація форми параметрів
        if not dep_permit:
            # якщо немає доступу до підрозділів
            userform.fields['dep'].widget.attrs.update({'disabled': True})
            userform.fields['section'].widget.attrs.update({'disabled': True})
            userform.fields['ntc_total'].widget.attrs.update({'disabled': True})
            userform.initial['employee'] = Employee.objects.filter(tab_nom=tab_nom).first()
        if can_change != 10: userform.fields['ntc_total'].widget.attrs.update({'disabled': True})
        if self.direction == '2':
            userform.fields['ntc_total'].widget.attrs.update({'hidden': 'hidden'})
            userform.fields['ntc_total'].label = ''

        if kind_rep == 1:
            if period:
                title = f"Підсумковий звіт по виконаним роботам {worker} за період" + \
                        f" з {work_date.strftime('%d.%m.%Y')} по {work_date_end.strftime('%d.%m.%Y')}"
            else:
                title = f"Підсумковий звіт по виконаним роботам {worker} за місяць"
        elif kind_rep == 2:
            if period:
                title = f"Звіт про розподіл витрат часу для виконаних робіт \
                        працівниками {worker} за період" + \
                        f" з {work_date.strftime('%d.%m.%Y')} по {work_date_end.strftime('%d.%m.%Y')}"
            else:
                title = f"Звіт про розподіл витрат часу для виконаних робіт \
                            працівниками {worker} за місяць"
        elif kind_rep == 3:
            title = f"Зведений звіт про розподіл питомих витрат часу (у %) для виконаних робіт {worker} за період"
            userform.initial['period'] = True
            userform.fields['period'].widget.attrs.update({'hidden': 'hidden'})
            userform.fields['period'].label = ''
            userform.fields['par_work_date'].label = 'Початок періоду'
            if work_date and work_date_end:
                title = title + f" з {work_date.strftime('%d.%m.%Y')} по {work_date_end.strftime('%d.%m.%Y')}"
        elif kind_rep == 4:
            title = f"Зведений звіт про розподіл витрат часу для виконаних робіт {worker} за період по замовленнях"
            userform.initial['period'] = True
            userform.fields['period'].widget.attrs.update({'hidden': 'hidden'})
            userform.fields['period'].label = ''
            userform.fields['par_work_date'].label = 'Початок періоду'
            if work_date and work_date_end:
                title = title + f" з {work_date.strftime('%d.%m.%Y')} по {work_date_end.strftime('%d.%m.%Y')}"

        if work_date:
            title = title.replace("місяць", f"{month_ua[work_date.month]} {work_date.strftime('%Y')} р.")

        if not params:
            params['page'] = 1
            params['limit'] = 0
        options_list = json.dumps(params, separators=(',', ':'))

        #  --- час формування звіту ----
        if start_time:
            end_time = time.time()
            execution_time_str += f"Заключні операції - {(end_time - start_time):.2f} сек\n"

        # --- Виведення сторінки  ----
        data = {"Title": title, "permit": request.user.get_all_permissions(),
                "work_date": work_date.strftime('%Y-%m-%d') if work_date else '',
                # "work_date_end": work_date_end.strftime('%Y-%m-%d') if work_date_end else '',
                "kind_rep": kind_rep,
                "form": userform,
                "title_fields": title_fields, "rows_list_value": report,
                "params": params, "options_list": options_list,
                "img_chart": '', "img_pie_chart": self.img_pie_chart, "total_table_html": self.total_table_html,
                "execution_time_str": execution_time_str,
                }

        #   --------- Друк Звіту (вигрузка у Word) ------
        if mode == 'print':
            if kind_rep == 1:
                return self.print_report(data)
            elif kind_rep == 2 or kind_rep == 3:
                return self.print_report_2_3(data)
            #  mode=='print' -  вигрузка у Word
        # ---- -  виведення діаграми ---
        elif mode == 'chart':
            data['img_chart'] = self.get_chart_3(data)

        # ----- Сторінка Звіту
        return render(request, self.report_template, context=data)

    # -------------------------------------

    @method_decorator(login_required)
    # def update_worktime(self, request, period, work_date, tab_nom, operation='', id_row=0, options=''):
    def update_worktime(self, request, operation='', id_row=0, options=''):
        """
        ---- Коригування/додавання/видалення рядка таблиці розподілу робочого часу НТЦ ----
        """
        # if isinstance(work_date, str): work_date = datetime.strptime(work_date, "%Y-%m-%d")
        rec = AccountWorkTime.objects.filter(pk=id_row).first()

        form_row = None
        nam_zak = ''
        #  ---- Збереження змін рядка ----
        if request.method == "POST":
            nam_zak = request.POST.get('nam_zak')
            form_row = WorkTime_rowForm(request.POST, instance=rec)  # , zakaz=request.POST.get('zakaz'), nam_zak=nam_zak)
            if not form_row.is_valid():
                # передаємо форму і список помилок у контекст шаблону
                return render(request, 'error_forms.html',
                              {'form': form_row, 'errors': form_row.errors})

        # ------------  видалення рядка -----------
        elif operation == 'del':
            #  визначення ключових параметрів по рядку БД
            if not rec:
                return HttpResponse(f"Рядок id={id_row} для видалення не знайдений !")
        else:
            return HttpResponse("Помилка запуску!")

        #  <<< --- Виконання операції в рамках транзакції
        try:
            with transaction.atomic():
                # ---- видалення рядка
                if operation == 'del':
                    rec.delete()
                #  --- збереження даних
                else:
                    #  збереження даних
                    new_rec = form_row.save()
                    #  К-ть документів, приведених до формату А4
                    new_rec.doc_kolA4 = float(new_rec.doc_format) * new_rec.doc_krform * new_rec.work_percent * 0.01
                    new_rec.save()
                    # визначаємо тип замовлення
                    type_zak = 'ЗЧ' if 'ЗапЧасти' in nam_zak else 'ОП'
                    #  додаємо Замовлення та Назву звмовлення в Історію замовлень, якщо його там немає
                    zak_hist = ZakazHistory.objects.filter(zakaz=new_rec.zakaz).first()
                    if zak_hist:
                        if zak_hist.name_zak[:100] != nam_zak[:100]:
                            #  коригування назви, якщо змінилася
                            zak_hist.name_zak = nam_zak[:100]
                            zak_hist.type_zak = type_zak
                            zak_hist.save()
                    else:
                        #  додавання в Історію
                        ZakazHistory.objects.create(zakaz=new_rec.zakaz, name_zak=nam_zak[:100], type_zak=type_zak)

        except Exception as e:
            return HttpResponse(f"Помилка {str(e)} операції {operation}з рядком id={id_row} !")
        # --------- >>>

        return HttpResponseRedirect(
            f"/worktime{self.direction if self.direction == '2' else ''}/{options}")

    @staticmethod
    def get_all_work_percent(request):
        """ ----------------------------------------------------------------------------------------------------------------
             Передача інфи про Загальний % готовності документу, суми нормооперацій та кількість аркушів для  запиту Ajax
         -------------------------------------------------------------------------------------------------------------------"""
        # Отримуємо параметри
        id_employee = request.POST.get('id_employee')
        id_zak = request.POST.get('id_zak')
        id_work_content = request.POST.get('id_work_content')
        id_doc_kind = request.POST.get('id_doc_kind')
        doc_designation = request.POST.get('doc_designation')
        id_row = request.POST.get('id_row')

        #  запит до БД по підрахунку
        queryset = AccountWorkTime.objects.filter(
            Q(employee=id_employee) & Q(zakaz=id_zak) & Q(work_content=id_work_content) & Q(doc_kind=id_doc_kind) &
            Q(doc_designation=doc_designation) & ~Q(id=id_row)
        ).all()
        sum_percent = queryset.aggregate(Sum("work_percent"))['work_percent__sum']
        sum_normoper = queryset.aggregate(Sum("kol_normoper"))['kol_normoper__sum']
        sum_dockrform = queryset.aggregate(Sum("doc_krform"))['doc_krform__sum']
        if not sum_percent: sum_percent = 0
        if not sum_normoper: sum_normoper = 0
        if not sum_dockrform: sum_dockrform = 0
        print('sum_percent= ', sum_percent)
        print('sum_normoper= ', sum_normoper)
        print('sum_dockrform= ', sum_dockrform)
        # Повернути відповідь в форматі JSON
        return JsonResponse({'percent_value': sum_percent, 'normoper_value': sum_normoper, 'dockrform_value': sum_dockrform})

    def print_report(self, data):
        """
             --- Друк Звіту 1:   виводимо в файл MS WORD
             data = {{ "Title": title,
                "work_date": work_date.strftime('%Y-%m-%d') if work_date else '',
                --- "tab_nom": tab_nom,
                "form": userform,
                "title_fields": title_fields, "rows_list_value": report }
        """
        docx_templ = self.docx_templ_1
        doc = Document(docx_templ)
        doc.styles['Normal'].font.size = Pt(11)
        # назва файлу
        docx_name = docx_templ.replace("dot", data['work_date'].replace('-', ''))
        # ----- Заповнюємо документ даними Звіту
        #  шукаємо параметр в тексті параграфа, якщо знайшли - замінюємо run.text для збереження формату
        # в одному параграфі - один параметр для заміни
        par1 = 'TITLE'
        for paragraph in doc.paragraphs:
            # шукаємо параметр на заміну
            repleced = False
            for run in paragraph.runs:
                if par1 in run.text:
                    run.text = run.text.replace(par1, data['Title'].replace('роботам ', 'роботам\n'))
                    repleced = True
                    break
            if repleced: break

        #  знаходимо таблицю в документі
        table = doc.tables[0]
        # записуємо в таблицю рядки Звіту
        for row in data['rows_list_value']:
            row_table = table.add_row()
            for i in range(0, 6):
                cell = row_table.cells[i]
                # Задаємо вирівнювання по висоті по центру
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = cell.paragraphs[0].add_run(str(row[i]))
                if 'Р А З О М' in row[1]: run.bold = True

        # ------------- підготовка документу для вигрузки
        f = BytesIO()
        doc.save(f)
        length = f.tell()
        f.seek(0)

        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_name
        response['Content-Length'] = length

        return response

    def print_report_2_3(self, data):
        """
             --- Друк Звітів  № 2 та 3:   виводимо в файл MS WORD
             data = {"Title": title,
                "work_date": work_date.strftime('%Y-%m-%d') if work_date else '',
                --- "tab_nom": tab_nom,
                "form": userform,
                "title_fields": title_fields, "rows_list_value": report,
                }

        """
        docx_templ = self.docx_templ_23
        doc = Document(docx_templ)
        doc.styles['Normal'].font.size = Pt(10)
        # назва файлу
        docx_name = docx_templ.replace("dot", data['work_date'].replace('-', ''))
        # ----- Заповнюємо документ даними Звіту
        #  шукаємо параметр в тексті параграфа, якщо знайшли - замінюємо run.text для збереження формату
        # в одному параграфі - один параметр для заміни
        par1 = 'TITLE'
        for paragraph in doc.paragraphs:
            # шукаємо параметр на заміну
            repleced = False
            for run in paragraph.runs:
                if par1 in run.text:
                    run.text = run.text.replace(par1, data['Title'].replace('робіт ', 'робіт\n'))
                    repleced = True
                    break
            if repleced: break

        # --- заміна заголовка в таблиці
        table = doc.tables[0]
        cells = table.rows[0].cells
        if data['kind_rep'] == 2:
            cells[1].paragraphs[0].runs[0].text = 'П.І.Б. (таб.номер),\nпосада'
        elif data['kind_rep'] == 3:
            cells[1].paragraphs[0].runs[0].text = 'Звітний період'

        #  знаходимо таблицю в документі
        table = doc.tables[1]
        # записуємо в таблицю рядки Звіту
        for row in data['rows_list_value']:
            row_table = table.add_row()
            for i in range(0, len(row)):
                cell = row_table.cells[i]
                # Задаємо вирівнювання по висоті по центру
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                run = cell.paragraphs[0].add_run('100' if row[i] == 100 else str(row[i]))
                # if 'Р А З О М' in row[1]: run.bold = True

        # ------------- підготовка документу для вигрузки
        f = BytesIO()
        doc.save(f)
        length = f.tell()
        f.seek(0)

        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_name
        response['Content-Length'] = length
        # ------>>>>>
        return response

    def get_report_01(self, templ_rep, queryset):
        """  ==============  Підсумковий звіт по виконаним роботам за місяць =========
            Шаблон Звіту в моделі - ReportWorkTimeTempl
        """
        report = []
        sum_time_total = sum_doc_kolA4_total = total_sum_time_docA4 = total_time_doc_A4 = 0
        total_sum_time_normoper = sum_time_normoper = sum_kol_normoper = total_kol_normoper = time_normoper = 0
        add_context = ('1', '1.3')
        for row_templ in templ_rep:
            row_templ.npp = row_templ.npp.strip()  # прибираємо пробіл зліва, який доданий для сортування моделі
            # --- по контексту робіт
            base_line = True
            queryset_row = queryset.filter(work_content=row_templ.work_content)
            context_name = row_templ.work_content.content
            if context_name.startswith("-"): context_name = context_name.replace("-", "", 1)
            #  -- по типу документу
            if row_templ.doc_type.id != 8:
                queryset_row = queryset_row.filter(doc_kind__doctype=row_templ.doc_type)
                context_name = row_templ.doc_type.name
                base_line = False
            # -- по виду документу
            if row_templ.doc_kind:
                queryset_row = queryset_row.filter(doc_kind=row_templ.doc_kind)
                context_name = row_templ.doc_kind.name
                base_line = False

            #  ---  сумуємо дані
            sum_queryset = queryset_row.aggregate(Sum("work_time"), Sum("doc_kolA4"))
            #  ---  сумуємо дані по нормоопераціям тільки по документу ТП(id=1026)
            if self.direction == '2':
                normoper_queryset = queryset_row.filter(doc_kind=1026)
                sum_normoper_queryset = normoper_queryset.aggregate(Sum("work_time"), Sum("kol_normoper"))

            #   витрачений час
            sum_time = sum_queryset['work_time__sum']
            if sum_time:
                if base_line: sum_time_total += sum_time
            else:
                sum_time = ''
            #   К-ть документів, приведених до формату А4
            sum_doc_kolA4 = sum_queryset['doc_kolA4__sum']
            if sum_doc_kolA4:
                sum_doc_kolA4 = round(sum_doc_kolA4, 2)
                if base_line: sum_doc_kolA4_total += sum_doc_kolA4
            else:
                sum_doc_kolA4 = ''
            #  Питома витрата часу на 1 документ, наведений до формату А4
            if sum_time and sum_doc_kolA4:
                time_doc_A4 = round(sum_time / sum_doc_kolA4, 2)
                if base_line:  total_sum_time_docA4 += sum_time
            else:
                time_doc_A4 = ''
            #  ------------------ для технологів - сума нормоперацій
            if self.direction == '2':
                # -- суми тільки по ТП
                if sum_normoper_queryset:
                    sum_kol_normoper = sum_normoper_queryset['kol_normoper__sum']
                    sum_time_normoper = sum_normoper_queryset['work_time__sum']
                else:
                    sum_kol_normoper = 0
                    sum_time_normoper = 0
                if sum_kol_normoper:
                    if base_line:
                        total_kol_normoper += sum_kol_normoper
                        total_sum_time_normoper += sum_time_normoper
                else:
                    sum_kol_normoper = ''
                #  Питома витрата часу на 1 нормооперацію
                if sum_time_normoper and sum_kol_normoper:
                    time_normoper = round(sum_time_normoper / sum_kol_normoper, 2)
                else:
                    time_normoper = ''
            #  ---------------
            #  -----------------  додаємо рядок до звіту ------------
            if row_templ.npp in add_context: context_name += ',\nв тому числі:'

            if self.direction == '2':
                # для технологів - сума нормоперацій
                report.append([row_templ.npp, context_name, sum_time, '', sum_doc_kolA4, time_doc_A4, sum_kol_normoper, time_normoper])
            else:
                report.append([row_templ.npp, context_name, sum_time, '', sum_doc_kolA4, time_doc_A4])

        #  --- рахуємо % по кожному рядку від загального часу
        for rep_row in report:
            if rep_row[2] and sum_time_total: rep_row[3] = round(rep_row[2] / sum_time_total * 100, 2)
        #  --- Totals sum
        if total_sum_time_docA4 and sum_doc_kolA4_total:
            total_time_doc_A4 = round(total_sum_time_docA4 / sum_doc_kolA4_total, 2)

        if self.direction == '2':
            # для технологів - сума нормоперацій
            if total_sum_time_normoper and total_kol_normoper:
                total_sum_time_normoper = round(total_sum_time_normoper / total_kol_normoper, 2)
            report.append(['', '\nР А З О М\n', sum_time_total, '', sum_doc_kolA4_total, total_time_doc_A4,
                           total_kol_normoper, total_sum_time_normoper])
        else:
            report.append(['', '\nР А З О М\n', sum_time_total, '', sum_doc_kolA4_total, total_time_doc_A4])

        title_fields = ['№\nп/п', 'Найменування виконаних робіт',
                        'Витрачена кількість часу на\nвиконання робіт, годин',
                        'Питома частка виконаних робіт\nвід витраченого часу, %',
                        'К-ть документів, наведених\nдо формату А4, од.',
                        'Питома витрата часу на 1 документ,\nнаведений до формату А4, година/А4'
                        ]
        if self.direction == '2':
            title_fields.extend(['Кількість нормооперацій', 'Питома витрата часу на\n1 нормооперацію, година/опер.'])
        #   ---
        return title_fields, report

    @staticmethod
    def get_report_02(templ_rep, queryset, fact_time_list):
        """ ======= Отчет о распределении затрат времени по выполненным работм
                    работниками бюро/отдела в  _____ 202__ г. ======  """
        time_rep_detail = ""
        work_date = queryset.last().work_date
        work_date = work_date.strftime('%Y-%m-%d')

        title_fields = []
        num = templ_rep.count()
        #   список унікальних id співробітників
        empl_list = []
        for e in queryset.values_list('employee', flat=True):
            if e not in empl_list:
                empl_list.append(e)

        npp = 0
        report = []
        # --- Пагінатор зроблено на клієнті, задаємо початкові параметри
        params = {'limit': 10, 'total_row': len(empl_list), 'page': 1}
        # <<<<< -----
        sum_time_total_all = sum_fact_time_all = 0  # загальний час по Звіту
        report_total_row = ['', 'Р А З О М:'] + [0] * num  # рядок з пудсумками по Звіту
        for empl_id in empl_list:
            npp += 1
            empl = Employee.objects.filter(pk=empl_id).first()
            report_row = [npp, f"{empl.full_name} (таб.№ {empl.tab_nom}),\n{empl.position}"] + [""] * num
            #  --- відбираємо інфу по співробітнику
            empl_queryset = queryset.filter(employee=empl)
            # --- визначаємо сумарні витрати часу співробітником по контекстам робіт та документам
            i = 1
            sum_time_total = 0
            for row_templ in templ_rep:
                # --- по контексту робіт
                i += 1
                base_line = True
                queryset_row = empl_queryset.filter(work_content=row_templ.work_content)
                #  -- по типу документу
                if row_templ.doc_type.id != 8:
                    queryset_row = queryset_row.filter(doc_kind__doctype=row_templ.doc_type)
                    base_line = False
                # -- по виду документу
                if row_templ.doc_kind:
                    queryset_row = queryset_row.filter(doc_kind=row_templ.doc_kind)
                    base_line = False

                #  ---  витрачений час
                sum_time = 0
                for s in queryset_row.values_list("work_time"):
                    sum_time += s[0]

                if sum_time:
                    if base_line: sum_time_total += sum_time
                    report_row[i] = sum_time  # додаємо в рядок звіту
                    report_total_row[i] += report_row[i]

            # ----- >>>>>
            report_row.append(sum_time_total)  # додаємо загальний час робіт
            # додаємо фактичний час по табелю
            fact_time = fact_time_list.get(empl.tab_nom, '')
            if fact_time: sum_fact_time_all += float(fact_time)
            report_row.append(fact_time)
            # загальний час по Звіту
            sum_time_total_all += sum_time_total
            # додаємо параметри для розшифровки рядка
            par_for_detal = json.dumps({"period": True, "work_date": work_date, "tab_nom": empl.tab_nom}, separators=(',', ':'))
            report_row.append(par_for_detal)
            report.append(report_row)  # додаємо рядок в звіт

        #  --- рахуємо суму годин по рядку з пудсумками по Звіту
        report_total_row.append(sum_time_total_all)  # додаємо загальний час робіт
        report_total_row.append(sum_fact_time_all)  # додаємо загальний фактичний час по табелю
        report.append(report_total_row)  # додаємо рядок в звіт
        #   ------------------ >>
        return title_fields, report, params, time_rep_detail

    def get_report_03(self, templ_rep, queryset, work_date, work_date_end):
        """ ===== Сводный отчет о распределении удельных затрат времени (в %) для выполненных работ
                  исполнителем/бюро/отделом за период ________ 202__ г. – ________ 202__ г.  """
        title_fields = []
        #  <<-- список місяців для звіту
        period_list = []
        date_start = work_date
        while date_start < work_date_end:
            date_end = date_start + relativedelta(months=1) - relativedelta(days=1)
            period_list.append((f"{month_ua[date_start.month]} {date_start.strftime('%Y')} р.", date_start, date_end))
            date_start = date_start + relativedelta(months=1)
        #  -- >>
        num = templ_rep.count()
        npp = 0
        report = []
        report_total_row = ['', 'Р А З О М:'] + [0] * num  # рядок з пудсумками по Звіту
        sum_time_total_all = 0  # загальний час по Звіту
        total_table_html_title = False

        for period_name, date_start, date_end in period_list:
            npp += 1
            report_row = [npp, period_name] + [''] * num
            #  --- відбираємо по місяцю
            period_queryset = queryset.filter(Q(work_date__gte=date_start) & Q(work_date__lte=date_end))
            # --- визначаємо сумарні витрати часу за період по контекстам робіт та документам
            i = 1
            sum_time_total = 0
            for row_templ in templ_rep:
                # --- по контексту робіт
                i += 1
                base_line = True
                queryset_row = period_queryset.filter(work_content=row_templ.work_content)
                #  -- по типу документу
                if row_templ.doc_type.id != 8:
                    queryset_row = queryset_row.filter(doc_kind__doctype=row_templ.doc_type)
                    base_line = False
                # -- по виду документу
                if row_templ.doc_kind:
                    queryset_row = queryset_row.filter(doc_kind=row_templ.doc_kind)
                    base_line = False
                #  ---  витрачений час
                sum_time = queryset_row.aggregate(Sum("work_time"))['work_time__sum']
                if sum_time:
                    if base_line: sum_time_total += sum_time  # загальний час по періоду
                    report_row[i] = sum_time  # додаємо в рядок звіту

            #  --- рахуємо % по кожному рядку від загального часу та пишемо в рядок звіту
            for i in range(2, len(report_row)):
                if report_row[i]:
                    report_total_row[i] += report_row[i]  # додаємо години в підсумки по Звіту
                    report_row[i] = round(report_row[i] / sum_time_total * 100, 2)
            sum_time_total_all += sum_time_total  # загальний час по Звіту
            report_row.append(sum_time_total)  # додаємо загальний час робіт
            report.append(report_row)  # додаємо рядок в звіт

            # ---------- підсумкова таблиця "Розподіл робочого часу" --------
            total_table, total_table_html = self.get_total_table(period_queryset)
            #  заголовки
            if not total_table_html_title:
                total_table_html_title = True
                self.total_table_html += [['&nbsp;Звітний&nbsp;період&nbsp;', ''] + [zak[1] for zak in total_table]]
            # ---------- Рядок підсумкової таблиці "Розподіл робочого часу" --------
            self.total_table_html += [[period_name] + total_table_html[0]]
            self.total_table_html += [[period_name] + total_table_html[1]]
            # --------------------------------

        #  --- рахуємо % по рядку з пудсумками по Звіту та пишемо в звіт
        for i in range(2, len(report_total_row)):
            if report_total_row[i]:
                report_total_row[i] = round(report_total_row[i] / sum_time_total_all * 100, 2)
            else:
                report_total_row[i] = ''
        report_total_row.append(sum_time_total_all)  # додаємо загальний час робіт
        report.append(report_total_row)  # додаємо рядок в звіт
        #   ------------------ >>
        return title_fields, report

    @staticmethod
    def get_report_04(templ_rep, queryset, options):
        """ ===== Сводный отчет о распределении затрат времени для выполненных работ
                  исполнителем/бюро/отделом ПО ЗАКАЗАМ за период """
        # -- список унікальних замовленнь для звіту
        list_zak = get_list_zak(9)  # довідник замовлень ()
        dict_zak = {zak[0]: zak[1].replace('&nbsp;', ' ') for zak in list_zak}
        rep_zak = queryset.order_by("zakaz").values_list("zakaz", flat=True)  # список замовлень у звіті
        title_fields = []
        zakaz_list = []
        for zak in list_zak:
            if zak[0] not in zakaz_list and zak[0] in rep_zak:
                zakaz_list.append(zak[0])

        num = templ_rep.count()
        npp = 0
        report = []
        report_total_row = ['', 'Р А З О М по звіту:'] + [0] * num  # рядок з пудсумками по Звіту
        sum_time_total_all = 0  # загальний час по Звіту
        # --- Пагінатор зроблено на клієнті, задаємо початкові параметри
        params = {'limit': 12, 'total_row': len(zakaz_list), 'page': 1, "main_field": "zakaz"}

        for zakaz in zakaz_list:
            npp += 1
            report_row = [npp, f"{zakaz} {dict_zak.get(zakaz, '???')}"] + [''] * num
            #  --- відбираємо по замовленню
            zakaz_queryset = queryset.filter(zakaz=zakaz)
            # --- визначаємо сумарні витрати часу по замовленню по контекстам робіт та документам
            i = 1
            sum_time_total = 0
            for row_templ in templ_rep:
                # --- по контексту робіт
                i += 1
                base_line = True
                queryset_row = zakaz_queryset.filter(work_content=row_templ.work_content)
                #  -- по типу документу
                if row_templ.doc_type.id != 8:
                    queryset_row = queryset_row.filter(doc_kind__doctype=row_templ.doc_type)
                    base_line = False
                # -- по виду документу
                if row_templ.doc_kind:
                    queryset_row = queryset_row.filter(doc_kind=row_templ.doc_kind)
                    base_line = False
                #  ---  витрачений час
                sum_time = queryset_row.aggregate(Sum("work_time"))['work_time__sum']
                if sum_time:
                    if base_line: sum_time_total += sum_time  # загальний час по періоду
                    report_row[i] = sum_time  # додаємо в рядок звіту

            #  --- рахуємо години по кожному рядку від загального часу та пишемо в рядок звіту
            for i in range(2, len(report_row)):
                if report_row[i]:
                    report_total_row[i] += report_row[i]  # додаємо години в підсумки по Звіту
                    #  О.Ю. 2023-10-16  - прийнято рішення замість %%  видавати в звіт години
                    # report_row[i] = round(report_row[i] / sum_time_total * 100, 2)
            sum_time_total_all += sum_time_total  # загальний час по Звіту
            report_row.append(sum_time_total)  # додаємо загальний час робіт
            report.append(report_row)  # додаємо рядок в звіт

        #  --- рахуємо суму годин по рядку з пудсумками по Звіту
        report_total_row.append(sum_time_total_all)  # додаємо загальний час робіт
        report.append(report_total_row)  # додаємо рядок в звіт
        #   ------------------ >>
        return title_fields, report, params

    @staticmethod
    def get_total_table(queryset):
        """ ======= Формування даних Розподілу робочого часу для Кругової діаграми та Підсумкової таблиці
        """
        # --- З довідника заказів беремо закази, які починаються з "ПР00", готуємо таблицю розподілу
        prefix_zak = "ПР00"
        prefix_zak_0 = f"{prefix_zak}00"
        total_table = [['ПР0000', 'Роботи по відкритим замовленням', 0]]
        total_table += [[zak[0], zak[1], 0] for zak in get_list_zak(9) if zak[0].startswith(prefix_zak)]
        # -- рахуємо суми по кожному замовленню з total_table
        sum_time_total = 0  # загальна сума витраченого часу
        for zak in total_table:
            if zak[0] == prefix_zak_0:  # роботи по відкритим замовленням
                res = queryset.exclude(zakaz__startswith=prefix_zak).aggregate(Sum("work_time"))
                zak[2] = res['work_time__sum'] if res['work_time__sum'] else 0
            else:  # роботи без відкритого замовлення
                res = queryset.filter(zakaz=zak[0]).aggregate(Sum("work_time"))
                zak[2] = res['work_time__sum'] if res['work_time__sum'] else 0
            sum_time_total += zak[2]

        # ------- дані для Підсумкової таблиці
        total_table_html = []
        total_table_html.append(['години'] + [zak[2] for zak in total_table])
        total_table_html.append(
            ['%'] + [round(zak[2] / sum_time_total * 100, 2) if sum_time_total > 0 else 0 for zak in total_table])

        return total_table, total_table_html

    def get_chart_12(self, queryset):
        """ ===================== Кругова діаграма в кінці Звіту - Розподіл робочого часу ================
        """
        # --- формуємо дані по queryset та готуємо таблицю розподілу та Діаграму
        total_table, total_table_html = self.get_total_table(queryset)

        # Дані для кругової діаграми
        labels = [zak[1].replace('&nbsp;', ' ') for zak in total_table if zak[2]]
        sizes = [zak[2] for zak in total_table if zak[2]]
        colors = ['lightskyblue', 'lightcoral', 'lightgreen', 'lightgrey', 'yellow', 'lightpink']

        # Очищення попередньої діаграми
        plt.clf()  # або plt.cla()

        # Створення об'єкта figure з вказаною шириною та висотою
        fig, ax = plt.subplots(figsize=(6, 6))

        # Побудова кругової діаграми
        plt.pie(sizes, autopct='%1.0f', startangle=90, colors=colors)  # autopct='%1.0f%%'

        # Встановлення назви діаграми
        plt.title(f"Розподіл робочого часу {self.title_dir}, %%", pad=1)

        # Додаткові налаштування
        plt.axis('equal')  # Забезпечує рівні пропорції вісей для отримання кругової форми

        # Виведення таблиці з повними назвами (легенда)
        legend = plt.legend(labels, loc='upper center', bbox_to_anchor=(0.5, -0.01))
        plt.tight_layout()

        # Зберігаємо діаграму у буфер пам'яті
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)

        # Кодуємо графік у формат base64
        string = base64.b64encode(buf.read())
        self.img_pie_chart = urllib.parse.quote(string)

        # ------- Формуємо підсумкову таблицю Розподілу робочого часу для виведення поруч з діаграмою
        #  заголовки + дані
        self.total_table_html = [[''] + [zak[1] for zak in total_table]] + total_table_html

    def get_chart_3(self, data):
        report = data['rows_list_value']
        # Створюємо дані для графіка
        x = []
        y = []
        y1 = []
        work_kd = [2, 7, 16]  # робота з КТД: Розробка, Внесення змін, Перевірка та узгодження
        work_other = [8, 9, 10, 11, 12, 13, 14, 15, 16, 17]
        for row in report:
            x.append(row[1])
            val_y = val_y1 = 0
            for i in range(2, len(row)):
                if i in work_kd:
                    val_y += float(row[i]) if row[i] else 0
                elif i in work_other:
                    val_y1 += float(row[i]) if row[i] else 0
            y.append(val_y)
            y1.append(val_y1)

        # # Побудова графіка
        # plt.plot(x, y)
        # plt.xlabel('Період')
        # plt.ylabel('% питомих витрат часу')
        # plt.title('Зведений звіт про розподіл питомих витрат часу')

        # Побудова стовпчастої діаграми
        # Встановлюємо розміри фігури (ширина=800 пікселів, висота=600 пікселів) і dpi=100
        fig, ax = plt.subplots(figsize=(12, 6), dpi=100)
        bar_width = 0.2
        opacity = 0.8
        index = np.arange(len(x))

        # ax.bar(x, y)
        ax.bar(index, y, bar_width, alpha=opacity, color='b', label='робота з КТД')
        ax.bar(index + bar_width, y1, bar_width, alpha=opacity, color='g', label='інші роботи')

        # Налаштування осей і заголовка
        ax.set_xlabel('')  # Період
        ax.set_ylabel('% питомих витрат часу')
        ax.set_title(f"Зведений звіт про розподіл питомих витрат часу по {self.title_dir}")
        ax.set_xticks(index + bar_width)
        ax.set_xticklabels(x)
        ax.legend()

        # Ростягнути діаграму
        plt.xticks(rotation=45, ha='right')
        plt.subplots_adjust(left=0.08, right=0.95, bottom=0.19, top=0.92)  # Встановлюємо відступи по осі x та осі y
        plt.tight_layout()

        # Зберігаємо графік у буфер пам'яті
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=100)
        buf.seek(0)

        # Кодуємо графік у формат base64
        string = base64.b64encode(buf.read())
        uri = urllib.parse.quote(string)

        return uri


#  --------- Розподіл робочого часу НТЦ ----------
work_time_ntc = WorkTime(direction="1", work_template="work_time.html", report_template="work_time_report.html")
work_time_tech = WorkTime(direction="2", work_template="work_time_tech.html", report_template="work_time_report_tech.html")
#  ---------------------------


"""
# формування QR-code
    import qrcode

    # Створіть URL-рядок з параметрами
    url_with_params = "com/path?" + "&".join(f"{key}={value}" for key, value in params.items())

    # Створіть QR-код
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(url_with_params)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")

    # Збережіть або відправте картинку до клієнта
    img.save("static/qrcode.png")  # Збереження QR-коду в файл

"""




'''
1523	53@68-30	Гребеник Наталія Валентинівна 
1146	58@71-1  Суханов В'ячеслав Валерійович


# # ********************  формування Історії замовлень (РАЗОВО!!!) додати до процедури Друку для виконання ************
# zak_list = Zakaz.get_list_zak(9)
# for zak in zak_list:
#     #  додаємо Замовлення та Назву звмовлення в Історію замовлень, якщо його там немає
#     print(f"{zak[0]} -- {zak[1]}")
#     zak, created = ZakazHistory.objects.get_or_create(zakaz=zak[0], name_zak=zak[1][:100])
# return HttpResponse("Історія замовлень сформована!")
# # *********************************************************************************


        # <<<<< -----
        for empl_id in empl_list:
            empl = Employee.objects.filter(pk=empl_id).first()
            npp += 1
            report_row = [npp, f"{empl.full_name} (таб.№ {empl.tab_nom}),\n{empl.position}"] + [""] * num
            #  --- відбираємо інфу по співробітнику
            empl_queryset = queryset.filter(employee=empl)
            # --- визначаємо сумарні витрати часу співробітником по контекстам робіт та документам
            i = 1
            sum_time_total = 0
            for row_templ in templ_rep:
                # --- по контексту робіт
                i += 1
                base_line = True
                queryset_row = empl_queryset.filter(work_content=row_templ.work_content)
                #  -- по типу документу
                if row_templ.doc_type.id != 8:
                    queryset_row = queryset_row.filter(doc_kind__doctype=row_templ.doc_type)
                    base_line = False
                # -- по виду документу
                if row_templ.doc_kind:
                    queryset_row = queryset_row.filter(doc_kind=row_templ.doc_kind)
                    base_line = False
                #  ---  витрачений час
                sum_time = queryset_row.aggregate(Sum("work_time"))['work_time__sum']
                if sum_time:
                    if base_line: sum_time_total += sum_time
                    report_row[i] = sum_time  # додаємо в рядок звіту
        # ----- >>>>>



'''
